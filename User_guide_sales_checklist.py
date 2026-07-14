from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_user_guide():
    doc = Document()
    
    # Title
    title = doc.add_heading('Sales Appointment Capture — User Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Version
    version = doc.add_paragraph('Version 1.6.2')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Helper function for headings
    def add_heading(text, level=1):
        h = doc.add_heading(text, level)
        return h

    # Helper function for paragraphs
    def add_para(text, style=None):
        p = doc.add_paragraph(text)
        if style:
            p.style = style
        return p

    # Helper for tables
    def add_table(headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # Bold headers
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        doc.add_paragraph() # Spacer

    # Content
    add_heading('1. Overview', 1)
    add_para('Sales Appointment Capture is an offline-capable Progressive Web App (PWA) for capturing client appointments in the field. It runs entirely on your device — no data is uploaded to any server.')
    
    add_para('What it does:', style='List Bullet')
    add_para('Capture client details, property addresses, and contact information', style='List Bullet')
    add_para('Complete Expression of Interest (EOI) and/or Irrevocable Authority (IA) forms', style='List Bullet')
    add_para('Take photos of client ID documents with your device camera', style='List Bullet')
    add_para('Capture client signatures directly on screen', style='List Bullet')
    add_para('Generate a compiled PDF of the entire appointment, plus a ZIP of separated documents for filing', style='List Bullet')
    add_para('Share the PDF and ZIP via email or your device\'s share sheet', style='List Bullet')
    
    add_para('Key fact: Everything stays on your device. PDFs are built in your browser. No cloud uploads, no server storage.', style='Quote')

    add_heading('2. Getting Started', 1)
    add_heading('Installing the App', 2)
    add_para('You only need to do this once per device:')
    add_para('Open the app URL in your browser while connected to the internet.', style='List Number')
    add_para('Wait for the page to load fully.', style='List Number')
    add_para('iPhone / iPad: Tap Safari\'s Share button → Add to Home Screen.', style='List Number')
    add_para('Android: Tap Chrome\'s menu → Add to Home screen / Install app.', style='List Number')
    add_para('Open the home-screen app once before leaving the office to confirm it loads.', style='List Number')
    add_para('After the first online visit, the app works fully offline — fill forms, attach photos, capture signatures, and generate PDFs without internet.')

    add_heading('Interface Tour', 2)
    add_para('The app has three main areas:')
    add_para('Header bar: App name, Save Draft, Load Draft, Load Test Data, Settings (gear icon), and version number.', style='List Bullet')
    add_para('Main content (left): The appointment form, divided into 6 numbered sections. Section 1 is always visible; Sections 2 (EOI) and 3 (IA) appear when you tick their checkboxes. The Appointment Summary Card at the top shows your progress at a glance.', style='List Bullet')
    add_para('Preview panel (right): A live PDF preview that updates when you click Refresh Preview or Generate PDF. Shows page count and navigation buttons.', style='List Bullet')
    add_para('Footer action bar: Generate PDF, Download PDF, Download Package, Share PDF, Save Draft, and New Appointment buttons. A live filename preview shows what the output will be named.', style='List Bullet')

    add_heading('3. Step-by-Step Appointment Flow', 1)
    add_heading('Section 1 — Appointment & Client Information', 2)
    
    s1_headers = ['Field', 'What to enter']
    s1_rows = [
        ['Date', 'Appointment date (auto-filled to today). Format: DD/MM/YYYY.'],
        ['Team Member', 'Your name. Free-text or select from a dropdown (configured in Settings).'],
        ['Residential Address', 'Client\'s current residential address.'],
        ['Property Sale Address', 'Address of the property being sold.'],
        ['Client 1 Name', 'Primary client\'s full name (required).'],
        ['Client 1 Phone', 'Primary client\'s contact number.'],
        ['Client 1 Email', 'Primary client\'s email address.'],
        ['Client 2 Name', 'Second client\'s full name (optional).'],
        ['Client 2 Phone', 'Second client\'s contact number.'],
        ['Client 2 Email', 'Second client\'s email address.']
    ]
    add_table(s1_headers, s1_rows)
    
    add_para('Checkboxes:')
    add_para('Include Expression of Interest — Tick to add EOI forms to the appointment.', style='List Bullet')
    add_para('Include Irrevocable Authority — Tick to add an IA form to the appointment.', style='List Bullet')
    add_para('When either checkbox is ticked, the corresponding section appears below with its fields.')

    add_heading('Section 2 — Expression of Interest (EOI)', 2)
    add_para('Visible when Include Expression of Interest is ticked.')
    add_para('Template: Choose Standard or La Vida Homes (the La Vida option reveals extra branded fields).')
    add_para('EOI Detail Overrides: Tick Show overrides to manually edit the client name, mobile, email, address, sale address, date, and staff member shown on the EOI form. When unticked, values are copied from Section 1 automatically.')
    
    add_para('Ownership:')
    add_para('Sole Owner — One client owns the property.', style='List Bullet')
    add_para('Joint Tenants 50/50 — Both clients own equally.', style='List Bullet')
    add_para('Tenants in Common — Enter custom percentage shares.', style='List Bullet')

    add_para('Sale Details:')
    s2_headers = ['Field', 'Notes']
    s2_rows = [
        ['Sale Type', 'Land, House, or House and Land.'],
        ['Land Price', 'Automatically formatted with $ and commas.'],
        ['House Price', 'Automatically formatted with $ and commas.'],
        ['H/L Total', 'Auto-calculated: Land + House.'],
        ['Finance %', 'Dropdown: 10% through 90%.'],
        ['Next Appointment Date', 'Date picker.'],
        ['Next Appointment Time', 'Dropdown: 8:00 AM – 6:00 PM in 30-minute steps.']
    ]
    add_table(s2_headers, s2_rows)

    add_para('La Vida Homes (visible when template = "La Vida Homes"):')
    add_para('Select a Finance Broker and Conveyancer from configured dropdowns, or choose "Other" to enter details manually.', style='List Bullet')
    
    add_para('Completion:')
    add_para('Branch Location — Brisbane or Perth (configurable in Settings).', style='List Bullet')
    add_para('Comments — Any extra notes.', style='List Bullet')

    add_heading('Section 3 — Irrevocable Authority (IA)', 2)
    add_para('Visible when Include Irrevocable Authority is ticked.')
    add_para('Form: Choose Perth IA Form or Brisbane IA Form.')
    add_para('IA Overrides: Tick Show overrides to manually edit client names, address, property, and date on the IA form. When unticked, values are copied from Section 1.')
    add_para('You can also click Copy EOI details to IA to pull values from the EOI section in one click.')
    
    add_para('Always-visible fields:')
    s3_headers = ['Field', 'Notes']
    s3_rows = [
        ['Authority Amount', 'Defaults to $10,000 (configurable in Settings).'],
        ['Solicitor / Conveyancer', 'Free-text or dropdown (configured in Settings).']
    ]
    add_table(s3_headers, s3_rows)
    add_para('Signature application: Tick Apply Sig 1 and/or Apply Sig 2 to stamp the captured signatures onto the IA form.')

    add_heading('Section 4 — Client ID Documents', 2)
    add_para('Four standard photo slots:')
    add_para('Client 1 ID Front', style='List Number')
    add_para('Client 1 ID Back', style='List Number')
    add_para('Client 2 ID Front', style='List Number')
    add_para('Client 2 ID Back', style='List Number')
    
    add_para('How to take a photo: Tap the relevant photo box. On mobile, your camera opens. On desktop, a file picker opens. Accepted formats: any image (JPEG, PNG, etc.).')
    add_para('Each photo slot has:')
    add_para('Rotate (↻) — Rotates the image 90° clockwise. Rotation is preserved in the PDF.', style='List Bullet')
    add_para('Remove (✕) — Removes the photo (asks for confirmation).', style='List Bullet')
    
    add_para('Additional Documents: Use the dropdown to add up to 6 extra document slots (Medicare Card, Passport, Driver Licence, Utility Bill, Bank Statement, Rates Notice, or Other). Each slot lets you assign the document to Client 1 or Client 2, choose the document type, and add a description for "Other" types.')
    add_para('Compression: The Compress for smaller PDF checkbox (ticked by default) reduces photo quality to keep PDFs smaller. Untick if you need maximum photo clarity.')

    add_heading('Section 5 — Signatures & Notes', 2)
    add_para('Signature 1: Draw here using a finger (mobile) or mouse (desktop). Label shows Client 1\'s name.')
    add_para('Signature 2: Optional second signature. Label shows Client 2\'s name.')
    add_para('Each signature canvas has a Clear button and a status badge showing Required, Optional, or Captured.')
    add_para('Notes: Free-text area for appointment notes, exceptions, or next steps.')

    add_heading('Section 6 — Pre-departure Checklist', 2)
    add_para('Before leaving the appointment, confirm:')
    add_para('Reservation Form is fully completed and signed by the client', style='List Bullet')
    add_para('Email addresses have been verified and double-checked', style='List Bullet')
    add_para('IA form details are completed and the correct Perth/Brisbane template is selected, if required', style='List Bullet')
    add_para('Clear photos of all client ID documents have been taken', style='List Bullet')

    add_heading('4. Auto-Fill & Smart Features', 1)
    add_para('Date Auto-Population: The appointment date is set to today automatically when the app opens. It formats as DD/MM/YYYY.')
    add_para('Copy EOI to IA: The Copy EOI details to IA button (in Section 3) copies client names, address, property address, and date from the EOI section into the IA override fields. This saves re-typing when both forms are needed.')
    add_para('Price Auto-Calculation: In the EOI section, the H/L Total field updates automatically as you type Land Price and House Price. Both price fields auto-format with $ and commas when you move to the next field.')
    add_para('Live Filename Preview: The footer bar shows a live preview of the output PDF filename: Sales Appointment - {date} - {client names} - {staff name}.pdf. It updates as you type.')
    add_para('Summary Card Click-to-Scroll: Click any item in the Appointment Summary Card to jump directly to that field in the form.')
    add_para('Progress Indicators: The summary card shows colour-coded indicators:')
    add_para('Grey circle — field is empty', style='List Bullet')
    add_para('Green tick — field is complete', style='List Bullet')
    add_para('Orange badge — items still remaining', style='List Bullet')
    add_para('Green badge — "Ready for PDF"', style='List Bullet')

    add_heading('5. Working with Photos', 1)
    add_heading('Taking Photos', 2)
    add_para('Tap the photo box you want to fill.', style='List Number')
    add_para('Your device camera opens (mobile) or a file picker appears (desktop).', style='List Number')
    add_para('Take the photo or select the file.', style='List Number')
    add_para('The photo appears as a thumbnail preview in the box.', style='List Number')

    add_heading('Rotating Photos', 2)
    add_para('If a photo is sideways or upside-down, tap the Rotate (↻) button. Each tap rotates 90° clockwise. Rotation is applied in the PDF.')

    add_heading('Removing Photos', 2)
    add_para('Tap the ✕ button. A confirmation dialog appears to prevent accidental removal.')

    add_heading('Additional Documents', 2)
    add_para('Need more than 4 photos? Select a number (1–6) from the Additional Documents dropdown at the bottom of Section 4. New photo slots appear. Choose the client, document type, and an optional description.')

    add_heading('Photo Compression', 2)
    add_para('The Compress for smaller PDF checkbox controls photo quality in the output:')
    add_para('Ticked (default): 2× render scale, 78% JPEG quality. Smaller files, good for email.', style='List Bullet')
    add_para('Unticked: 3× render scale, 92% JPEG quality. Larger files, maximum detail.', style='List Bullet')

    add_heading('6. Sharing & Sending', 1)
    add_heading('Step 1: Generate the PDF', 2)
    add_para('Click Generate PDF in the header or footer bar. The app validates all required fields, builds the PDF, and shows a preview in the right panel. Use Prev / Next to flip through pages. The status bar shows the file size.')

    add_heading('Step 2: Choose your output method', 2)
    share_headers = ['Button', 'What it does']
    share_rows = [
        ['Download PDF', 'Downloads the compiled PDF to your device.'],
        ['Download Package', 'Downloads the compiled PDF plus a ZIP containing individual separated PDFs — one for the EOI, one for the IA, and one per photo.'],
        ['Share PDF', 'Opens your device\'s share sheet (if supported) or falls back to downloading the files and opening a pre-filled email.']
    ]
    add_table(share_headers, share_rows)

    add_heading('Sharing via Email', 2)
    add_para('When you tap Share PDF:')
    add_para('If your device supports file sharing, both the compiled PDF and the ZIP of separated documents are attached to the share sheet. You can send them via Mail, Messages, AirDrop, or any other app.', style='List Number')
    add_para('If file sharing is not supported (e.g., Firefox on desktop, or HTTP connections), the app downloads both files and opens a pre-filled email draft addressed to the configured recipients with the subject and body already filled in. Attach the downloaded files manually before sending.', style='List Number')
    add_para('The email is pre-addressed to the team\'s processing inbox. The subject and body include the client names, property address, date, and staff member name.')

    add_heading('Web Share Compatibility', 2)
    compat_headers = ['Platform', 'Browser', 'File sharing works?']
    compat_rows = [
        ['iPhone / iPad', 'Safari', 'Yes'],
        ['Android', 'Chrome', 'Yes'],
        ['Windows Desktop', 'Chrome / Edge (HTTPS)', 'Yes'],
        ['Windows Desktop', 'Firefox', 'No (falls back to email)'],
        ['Mac Desktop', 'Safari', 'No (falls back to email)'],
        ['Mac Desktop', 'Chrome (HTTPS)', 'Yes']
    ]
    add_table(compat_headers, compat_rows)

    add_heading('7. Drafts & Reset', 1)
    add_heading('Saving a Draft', 2)
    add_para('At any point, click Save Draft (header or footer). All form fields, photos, and signatures are saved to your device\'s browser storage. A confirmation toast appears.')
    add_para('Note: Only one draft exists at a time. Saving overwrites the previous draft.')

    add_heading('Loading a Draft', 2)
    add_para('Click Load Draft (header). The form restores to its last saved state. If no draft exists, a message appears: "No saved draft found on this device."')
    add_para('Important: Drafts are stored on the specific device and browser you used. A draft saved on your phone is not available on your laptop.')

    add_heading('Load Test Data', 2)
    add_para('Click Load Test Data (header, beaker icon) to populate the form with sample data for testing or demonstration. A confirmation dialog appears first.')

    add_heading('Starting a New Appointment', 2)
    add_para('Click New Appointment (footer, bin icon) to clear all fields, photos, and signatures. A confirmation dialog appears. The date resets to today.')

    add_heading('8. Settings', 1)
    add_para('Tap the gear icon (⚙) in the header to open Settings.')
    
    add_heading('Unlocking Settings', 2)
    add_para('Enter the admin PIN and tap Unlock. The unlock state lasts until you close the browser tab.')

    add_heading('Settings Tabs', 2)
    add_para('Dropdowns:')
    add_para('Staff Members — Switch between free-text input and a preset dropdown. Add or remove staff names.', style='List Bullet')
    add_para('Branch Locations — Add or remove branch options (default: Perth, Brisbane).', style='List Bullet')
    add_para('Solicitors / Conveyancers — Switch between free-text and dropdown. Add or remove names.', style='List Bullet')
    add_para('EOI Templates — Add or remove template options (default: Standard, La Vida Homes).', style='List Bullet')
    add_para('Additional Document Types — Customise the document type dropdown (default: Medicare Card, Passport, Driver Licence, Utility Bill, Bank Statement, Rates Notice, Other).', style='List Bullet')
    
    add_para('PDF Defaults:')
    add_para('Default authority amount (e.g. $10,000)', style='List Bullet')
    add_para('Default finance percentage', style='List Bullet')
    add_para('Default branch location', style='List Bullet')
    add_para('Compress photos by default (on/off)', style='List Bullet')
    add_para('Apply Signature 1 / Signature 2 to IA by default (on/off)', style='List Bullet')
    
    add_para('La Vida:')
    add_para('Add, edit, or remove Finance Broker contacts (name, email, phone).', style='List Bullet')
    add_para('Add, edit, or remove Settlement / Conveyancer contacts.', style='List Bullet')
    
    add_para('Company:')
    add_para('Display-only — company details are built into the PDF templates.', style='List Bullet')

    add_heading('Import / Export Settings', 2)
    add_para('Use Export Settings to download your configuration as a JSON file. Use Import Settings to load a configuration from another device. This is useful for setting up multiple staff devices with the same dropdowns and defaults.')

    add_heading('9. Tips & Troubleshooting', 1)
    add_heading('Offline Use', 2)
    add_para('The app works fully offline after the first online visit. You do not need internet to fill forms, take photos, capture signatures, or generate PDFs. PDFs are built entirely on your device. Large photo sets may take a few seconds.')

    add_heading('Photo Storage', 2)
    add_para('Photos are stored as part of your draft in browser local storage, which has a limit of roughly 5–10 MB per website. If you take many high-resolution photos, a draft save may fail with a storage warning. Reduce the number of photos or enable compression.')

    add_heading('Getting Updates', 2)
    add_para('When a new version of the app is released, open it once while online. The new version caches automatically. The version number is displayed in the header and at the bottom of every generated PDF page.')

    add_heading('PDF Not Generating?', 2)
    add_para('Check that required fields are filled: Client 1 Name, Team Member, and Date.', style='List Bullet')
    add_para('Check that at least one form (EOI or IA) is ticked to include.', style='List Bullet')
    add_para('Any field with a red border and error message needs attention.', style='List Bullet')

    add_heading('Photos Not Appearing in PDF?', 2)
    add_para('Ensure photos are attached (thumbnail visible in the photo slot).', style='List Bullet')
    add_para('Check that the photo was not accidentally removed (the slot should show a thumbnail, not "Tap to attach").', style='List Bullet')

    add_heading('Draft Won\'t Load?', 2)
    add_para('Drafts are device- and browser-specific. A draft saved in Safari won\'t appear in Chrome.', style='List Bullet')
    add_para('Clearing your browser data removes the draft. Export important settings before clearing.', style='List Bullet')

    doc.save('Sales_Appointment_Capture_User_Guide.docx')
    print("Document created successfully!")

if __name__ == "__main__":
    try:
        create_user_guide()
    except ImportError:
        print("Please install python-docx first: pip install python-docx")