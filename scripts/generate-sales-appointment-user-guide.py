"""Build the ASG Sales Appointment Capture staff guide from canonical Markdown."""

from __future__ import annotations

import argparse
import hashlib
import re
import subprocess
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/user-guides/source/SALES_APPOINTMENT_CAPTURE_USER_GUIDE.md"
OUT_DIR = ROOT / "docs/user-guides"
DOCX_PATH = OUT_DIR / "ASG_Sales_Appointment_Capture_User_Guide.docx"
PDF_PATH = OUT_DIR / "ASG_Sales_Appointment_Capture_User_Guide.pdf"
SCREENSHOTS_JSON = ROOT / "docs/user-guides/screenshots.json"
SCREENSHOT_DIR = ROOT / "docs/user-guides/screenshots"
LIBREOFFICE: Path | None = None
LIBREOFFICE_PROFILE: Path | None = None
METADATA: dict[str, str] = {}
NAVY = "071B33"
GOLD = "C6A14A"
PALE_GOLD = "FBF5E7"
PALE_BLUE = "F3F6FA"
TEXT = RGBColor(13, 37, 63)
MUTED = RGBColor(88, 105, 126)


def shade(cell, colour: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    fill = OxmlElement("w:shd")
    fill.set(qn("w:fill"), colour)
    properties.append(fill)


def set_cell_border(cell, *, left: str | None = None, bottom: str | None = None) -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge, colour, size in (("left", left, "18"), ("bottom", bottom, "4")):
        if colour:
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"), "single")
            tag.set(qn("w:sz"), size)
            tag.set(qn("w:color"), colour)
            borders.append(tag)


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for style_name, size, colour in (
        ("Title", 32, RGBColor(255, 255, 255)),
        ("Heading 1", 20, TEXT),
        ("Heading 2", 12, TEXT),
        ("Heading 3", 10, TEXT),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = colour
        style.paragraph_format.keep_with_next = True


def add_cover(document: Document) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Mm(210)
    row = table.rows[0]
    row.height = Mm(296)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = row.cells[0]
    cell.width = Mm(210)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade(cell, NAVY)
    margins = cell._tc.get_or_add_tcPr()
    cell_mar = OxmlElement("w:tcMar")
    for edge, value in (("top", "900"), ("left", "1100"), ("bottom", "900"), ("right", "1100")):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), value)
        el.set(qn("w:type"), "dxa")
        cell_mar.append(el)
    margins.append(cell_mar)

    logo = cell.paragraphs[0]
    logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo.add_run().add_picture(str(ROOT / "icons/asg_logo.png"), width=Cm(7.2))
    label = cell.add_paragraph("ASG INTERNAL STAFF RESOURCE")
    label.paragraph_format.space_before = Pt(38)
    label.paragraph_format.space_after = Pt(10)
    run = label.runs[0]
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(GOLD)
    run.font.letter_spacing = Pt(1.4)
    title = cell.add_paragraph("Sales Appointment\nCapture")
    title.style = document.styles["Title"]
    title.paragraph_format.space_after = Pt(14)
    subtitle = cell.add_paragraph("Staff User Guide")
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor.from_string(GOLD)
    purpose = cell.add_paragraph("A practical guide to completing, reviewing and handing over in-person and Zoom sales appointments.")
    purpose.paragraph_format.space_before = Pt(18)
    purpose.paragraph_format.space_after = Pt(45)
    purpose.runs[0].font.size = Pt(12)
    purpose.runs[0].font.color.rgb = RGBColor(220, 227, 235)
    metadata = cell.add_paragraph(
        "\n".join(
            (
                f"Application version: {METADATA['Application version']}",
                f"Guide version: {METADATA['Guide version']}",
                f"Generated: {METADATA['Generated']}",
                f"Git branch: {METADATA['Git branch']}",
                f"Source commit: {METADATA['Source commit']}",
            )
        )
    )
    for run in metadata.runs:
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(178, 190, 204)


def configure_content_section(document: Document):
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(17)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.text = "ASG  |  SALES APPOINTMENT CAPTURE"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(7.5)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(GOLD)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(
        "ASG Internal Staff Resource"
        f"  |  App {METADATA['Application version']}"
        f"  |  Guide {METADATA['Guide version']}  |  "
    )
    add_page_field(footer)
    for run in footer.runs:
        run.font.name = "Arial"
        run.font.size = Pt(7.5)
        run.font.color.rgb = MUTED
    return section


def add_rich_paragraph(document: Document, text: str, style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.name = "Arial"
    return paragraph


def add_callout(document: Document, lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, PALE_GOLD)
    set_cell_border(cell, left=GOLD)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        for part in re.split(r"(\*\*.+?\*\*)", line):
            run = paragraph.add_run(part[2:-2] if part.startswith("**") else part)
            run.bold = part.startswith("**")
            run.font.size = Pt(9)
            run.font.color.rgb = TEXT
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        if row_index == 0:
            set_repeat_table_header(row)
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.text = value.replace("**", "")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                shade(cell, NAVY)
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if column_index == 0:
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                        run.font.color.rgb = TEXT
                set_cell_border(cell, bottom="DDE3EA")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_image(document: Document, relative_path: str, caption: str) -> None:
    path = SOURCE.parent / relative_path
    with Image.open(path) as source_image:
        aspect = source_image.width / source_image.height
    width_cm = 15.5
    height_cm = width_cm / aspect
    if height_cm > 12.2:
        height_cm = 12.2
        width_cm = height_cm * aspect
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm), height=Cm(height_cm))
    paragraph.paragraph_format.space_after = Pt(2)
    caption_paragraph = document.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(4)
    caption_run = caption_paragraph.runs[0]
    caption_run.italic = True
    caption_run.font.size = Pt(7.5)
    caption_run.font.color.rgb = MUTED


def render_section(document: Document, heading: str, body: str, first: bool) -> None:
    if not first:
        document.add_page_break()
    title = document.add_heading(heading, level=1)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    rule_table = document.add_table(rows=1, cols=1)
    rule_cell = rule_table.cell(0, 0)
    rule_cell.text = ""
    shade(rule_cell, GOLD)
    rule_table.rows[0].height = Mm(1.2)
    rule_table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    document.add_paragraph().paragraph_format.space_after = Pt(0)

    lines = body.strip().splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line == "---":
            index += 1
            continue
        if line.startswith("### "):
            add_rich_paragraph(document, line[4:], "Heading 2")
        elif line.startswith("!["):
            match = re.match(r"!\[(.+?)\]\((.+?)\)", line)
            if match:
                add_image(document, match.group(2), match.group(1))
        elif line.startswith("> "):
            callout = []
            while index < len(lines) and lines[index].startswith(">"):
                callout.append(lines[index].lstrip("> ").rstrip())
                index += 1
            add_callout(document, callout)
            continue
        elif line.startswith("|"):
            rows = []
            while index < len(lines) and lines[index].startswith("|"):
                values = [value.strip() for value in lines[index].strip("|").split("|")]
                if not all(set(value) <= {"-", ":"} for value in values):
                    rows.append(values)
                index += 1
            add_markdown_table(document, rows)
            continue
        elif re.match(r"^\d+\. ", line):
            add_rich_paragraph(document, line)
        elif line.startswith("- [ ] "):
            add_rich_paragraph(document, "☐ " + line[6:])
        elif line.startswith("- "):
            add_rich_paragraph(document, line[2:], "List Bullet")
        else:
            add_rich_paragraph(document, line)
        index += 1


def parse_generated_metadata(markdown: str) -> dict[str, str]:
    start_marker = "<!-- docs-automation:metadata:start -->"
    end_marker = "<!-- docs-automation:metadata:end -->"
    if markdown.count(start_marker) != 1 or markdown.count(end_marker) != 1:
        raise RuntimeError("Expected exactly one generated metadata marker pair")
    start = markdown.index(start_marker)
    end = markdown.index(end_marker)
    if start >= end:
        raise RuntimeError("Generated metadata markers are reversed")
    lines = [
        line.strip()
        for line in markdown[start + len(start_marker):end].splitlines()
        if line.strip()
    ]
    expected = (
        "Application version",
        "Guide version",
        "Generated",
        "Git branch",
        "Source commit",
    )
    if len(lines) != len(expected):
        raise RuntimeError("Generated metadata block is malformed")
    metadata: dict[str, str] = {}
    for label, line in zip(expected, lines, strict=True):
        match = re.fullmatch(r"\*\*([^*]+):\*\*\s+(.+?)(?:<br>)?", line)
        if not match or match.group(1) != label or not match.group(2).strip():
            raise RuntimeError(f"Generated metadata field is malformed: {label}")
        metadata[label] = match.group(2).strip()
    return metadata


def normalise_docx(path: Path) -> None:
    temporary = path.with_suffix(".normalised.docx")
    with zipfile.ZipFile(path, "r") as source_zip, zipfile.ZipFile(
        temporary,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
    ) as output_zip:
        for name in sorted(source_zip.namelist()):
            original = source_zip.getinfo(name)
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = original.external_attr
            info.create_system = original.create_system
            output_zip.writestr(info, source_zip.read(name))
    temporary.replace(path)


def normalise_pdf(path: Path, generated: str, document_hash: str) -> None:
    generated_date = datetime.strptime(generated, "%d %B %Y")
    fixed_date = generated_date.strftime("D:%Y%m%d120000+08'00'")
    content = path.read_bytes()
    content = re.sub(
        rb"/CreationDate\s*\(D:[^)]*\)",
        f"/CreationDate({fixed_date})".encode("ascii"),
        content,
    )
    identifier = document_hash[:32].upper().encode("ascii")
    content = re.sub(
        rb"/ID\s*\[\s*<[0-9A-Fa-f]{32}>\s*<[0-9A-Fa-f]{32}>\s*\]",
        b"/ID [ <" + identifier + b">\n<" + identifier + b"> ]",
        content,
    )
    path.write_bytes(content)


def build_docx() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    global METADATA
    METADATA = parse_generated_metadata(markdown)
    sections = re.findall(r"^## (\d+\..+?)\n(.*?)(?=^## \d+\.|\Z)", markdown, flags=re.M | re.S)
    if len(sections) != 16:
        raise RuntimeError(f"Expected 16 numbered sections, found {len(sections)}")
    document = Document()
    configure_document(document)
    add_cover(document)
    configure_content_section(document)
    for index, (heading, body) in enumerate(sections):
        render_section(document, heading.strip(), body, first=index == 0)
    document.core_properties.title = "Sales Appointment Capture — Staff User Guide"
    document.core_properties.subject = "ASG internal staff operating guide"
    document.core_properties.author = "Amplify Solutions Group"
    document.core_properties.comments = (
        "Generated from canonical repository Markdown. "
        f"Application version {METADATA['Application version']}; "
        f"Guide version {METADATA['Guide version']}; "
        f"Generated {METADATA['Generated']}; "
        f"Git branch {METADATA['Git branch']}; "
        f"Source commit {METADATA['Source commit']}."
    )
    fixed_datetime = datetime.strptime(METADATA["Generated"], "%d %B %Y")
    document.core_properties.created = fixed_datetime
    document.core_properties.modified = fixed_datetime
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(DOCX_PATH)
    normalise_docx(DOCX_PATH)


def export_pdf() -> None:
    if LIBREOFFICE is None or LIBREOFFICE_PROFILE is None:
        raise RuntimeError("Validated LibreOffice executable and isolated profile are required")
    LIBREOFFICE_PROFILE.mkdir(parents=True, exist_ok=True)
    command = [
        str(LIBREOFFICE),
        "--headless",
        "--nologo",
        "--nodefault",
        "--nolockcheck",
        "--nofirststartwizard",
        f"-env:UserInstallation={LIBREOFFICE_PROFILE.as_uri()}",
        "--convert-to",
        "pdf:writer_pdf_Export",
        "--outdir",
        str(OUT_DIR),
        str(DOCX_PATH),
    ]
    result = subprocess.run(
        command,
        cwd=ROOT,
        capture_output=True,
        text=True,
        timeout=120,
        shell=False,
    )
    if result.returncode:
        raise RuntimeError(result.stderr or result.stdout)
    generated = OUT_DIR / f"{DOCX_PATH.stem}.pdf"
    if generated != PDF_PATH and generated.exists():
        generated.replace(PDF_PATH)
    if not PDF_PATH.exists():
        raise RuntimeError("LibreOffice did not produce the expected PDF")
    normalise_pdf(
        PDF_PATH,
        METADATA["Generated"],
        hashlib.sha256(DOCX_PATH.read_bytes()).hexdigest(),
    )


def configure_paths(arguments: argparse.Namespace) -> None:
    global SOURCE, OUT_DIR, DOCX_PATH, PDF_PATH
    global SCREENSHOTS_JSON, SCREENSHOT_DIR, LIBREOFFICE, LIBREOFFICE_PROFILE
    SOURCE = arguments.source.resolve()
    OUT_DIR = arguments.output_dir.resolve()
    DOCX_PATH = OUT_DIR / "ASG_Sales_Appointment_Capture_User_Guide.docx"
    PDF_PATH = OUT_DIR / "ASG_Sales_Appointment_Capture_User_Guide.pdf"
    SCREENSHOTS_JSON = arguments.screenshots_json.resolve()
    SCREENSHOT_DIR = arguments.screenshot_dir.resolve()
    LIBREOFFICE = arguments.libreoffice
    LIBREOFFICE_PROFILE = arguments.profile_dir.resolve()


def parse_arguments() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--screenshots-json", type=Path, required=True)
    parser.add_argument("--screenshot-dir", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--libreoffice", type=Path, required=True)
    parser.add_argument("--profile-dir", type=Path, required=True)
    return parser.parse_args()


if __name__ == "__main__":
    configure_paths(parse_arguments())
    build_docx()
    export_pdf()
    print(f"Created {DOCX_PATH.relative_to(ROOT)}")
    print(f"Created {PDF_PATH.relative_to(ROOT)}")
